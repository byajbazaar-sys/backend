import fitz
import docx2txt
import spacy
from pathlib import Path
from typing import Dict, List, Any, Tuple
import logging
import json
from datetime import datetime
import re
import base64
from pdfminer.high_level import extract_text as pdfminer_extract
from PIL import Image
from docx import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors"""
    pass


class DualPDFDocxParserWithLinksAndImages:
    """
    Advanced resume parser with:
    - DUAL PDF extraction (PyMuPDF + pdfminer.six)
    - HYPERLINK extraction (PDF and DOCX)
    - IMAGE extraction (PDF and DOCX) with Base64 encoding
    - Multiple information extraction methods
    
    Usage:
        parser = DualPDFDocxParserWithLinksAndImages()
        result = parser.parse("resume.pdf")
        parser.export_to_json(result, "output.json")
        parser.export_to_csv(result, "output.csv")
    """
    
    def __init__(self, model_name: str = "en_core_web_sm"):
        """Initialize parser with spaCy model"""
        try:
            model_path = os.getenv("SPACY_DATA_PATH", None)
            if model_path:
                nlp = spacy.load(model_path + "/en_core_web_sm")
            else:
                nlp = spacy.load("en_core_web_sm")
            logger.info(f"Loaded spaCy model: {model_name}")
        except OSError as e:
            logger.error(f"Failed to load spaCy model: {e}")
            raise ResumeParsingError(
                f"spaCy model '{model_name}' not found. "
                "Run: python -m spacy download en_core_web_sm"
            ) from e
    
    # ==================== File Extraction with Hyperlinks & Images ====================
    
    def extract_resume_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text, hyperlinks, and images from PDF or DOCX files
        
        Returns:
            Tuple of (merged_text, extraction_details)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ResumeParsingError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf_complete(str(file_path))
            elif extension == '.docx':
                return self._extract_from_docx_complete(str(file_path))
            else:
                raise ResumeParsingError(
                    f"Unsupported file type: {extension}. "
                    "Supported: .pdf, .docx"
                )
        except Exception as e:
            logger.error(f"Error extracting from {file_path}: {e}")
            raise ResumeParsingError(f"Failed to extract: {str(e)}") from e
    
    # ==================== PDF Complete Extraction ====================
    
    def _extract_from_pdf_complete(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text, hyperlinks, and images from PDF"""
        extraction_details = {}
        
        # Extract images from PDF with Base64 encoding
        logger.info("Extracting images from PDF...")
        images = self._extract_images_from_pdf(file_path)
        extraction_details["images"] = images
        logger.info(f"✅ Found {len(images)} image(s)")
        
        # Extract hyperlinks from PDF
        logger.info("Extracting hyperlinks from PDF...")
        hyperlinks = self._extract_hyperlinks_from_pdf(file_path)
        extraction_details["hyperlinks"] = hyperlinks
        logger.info(f"✅ Found {len(hyperlinks)} hyperlink(s)")
        
        # Extract text with PyMuPDF
        try:
            logger.info("Extracting text with PyMuPDF...")
            pymupdf_text = self._extract_text_with_pymupdf(file_path)
            extraction_details["pymupdf"] = pymupdf_text
            logger.info("✅ PyMuPDF extraction successful")
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
            extraction_details["pymupdf"] = ""
            pymupdf_text = ""
        
        # Extract text with pdfminer
        try:
            logger.info("Extracting text with pdfminer...")
            pdfminer_text = self._extract_text_with_pdfminer(file_path)
            extraction_details["pdfminer"] = pdfminer_text
            logger.info("✅ pdfminer extraction successful")
        except Exception as e:
            logger.warning(f"pdfminer extraction failed: {e}")
            extraction_details["pdfminer"] = ""
            pdfminer_text = ""
        
        # Merge text extractions
        merged_text = self._merge_pdf_extractions(pymupdf_text, pdfminer_text)
        
        if not merged_text.strip():
            raise ResumeParsingError("Both PDF extractors returned no text content")
        
        logger.info("PDF extraction (complete) successful")
        return merged_text, extraction_details
    
    def _extract_images_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract images from PDF with Base64 encoding
        
        Returns list of dicts with:
        - image_name: Filename
        - image_base64: Base64 encoded image data
        - image_format: Format (png, jpeg, etc)
        - page_number: Which page the image is on
        - extraction_method: "PyMuPDF"
        """
        images = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                try:
                    # Get images on this page
                    image_list = page.get_images()
                    
                    for img_index, img_ref in enumerate(image_list):
                        try:
                            xref = img_ref[0]
                            
                            # Extract image as binary data
                            img_data = doc.extract_image(xref)
                            image_bytes = img_data["image"]
                            image_ext = img_data["ext"]
                            
                            # Convert to Base64
                            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                            
                            images.append({
                                "image_name": f"page_{page_num}_image_{img_index + 1}.{image_ext}",
                                "image_base64": image_base64,
                                "image_format": image_ext,
                                "page_number": page_num,
                                "image_size": len(image_bytes),
                                "extraction_method": "PyMuPDF"
                            })
                            
                            logger.debug(f"Extracted image on page {page_num}: {len(image_bytes)} bytes")
                        
                        except Exception as e:
                            logger.warning(f"Error processing image on page {page_num}: {e}")
                            continue
                
                except Exception as e:
                    logger.warning(f"Error extracting images from page {page_num}: {e}")
                    continue
            
            doc.close()
            return images
        
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return []
    
    def _extract_hyperlinks_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract hyperlinks from PDF"""
        hyperlinks = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                try:
                    links = page.get_links()
                    
                    for link in links:
                        try:
                            if "uri" in link:
                                url = link["uri"]
                                link_rect = link.get("from")
                                
                                link_text = ""
                                if link_rect:
                                    link_text = page.get_textbox(link_rect).strip()
                                
                                if not link_text:
                                    words = page.get_text("words")
                                    for word in words:
                                        word_rect = fitz.Rect(word[:4])
                                        if link_rect and word_rect.intersects(link_rect):
                                            link_text = word[4]
                                            break
                                
                                hyperlinks.append({
                                    "link_text": link_text,
                                    "link_url": url,
                                    "page_number": page_num,
                                    "coordinates": str(link_rect) if link_rect else "N/A",
                                    "extraction_method": "PyMuPDF"
                                })
                                
                                logger.debug(f"Found link on page {page_num}: {link_text} -> {url}")
                        
                        except Exception as e:
                            logger.warning(f"Error processing link on page {page_num}: {e}")
                            continue
                
                except Exception as e:
                    logger.warning(f"Error extracting links from page {page_num}: {e}")
                    continue
            
            doc.close()
            return hyperlinks
        
        except Exception as e:
            logger.error(f"Hyperlink extraction failed: {e}")
            return []
    
    def _extract_text_with_pymupdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ''
            
            for page_num, page in enumerate(doc, 1):
                try:
                    page_text = page.get_text()
                    text += f"\n--- Page {page_num} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"PyMuPDF failed to extract page {page_num}: {e}")
                    continue
            
            doc.close()
            return text
        
        except Exception as e:
            raise ResumeParsingError(f"PyMuPDF extraction failed: {str(e)}") from e
    
    def _extract_text_with_pdfminer(self, file_path: str) -> str:
        """Extract text from PDF using pdfminer.six"""
        try:
            text = pdfminer_extract(file_path)
            return text
        except Exception as e:
            raise ResumeParsingError(f"pdfminer extraction failed: {str(e)}") from e
    
    def _merge_pdf_extractions(self, pymupdf_text: str, pdfminer_text: str) -> str:
        """Merge PyMuPDF and pdfminer extractions"""
        if not pdfminer_text.strip():
            return pymupdf_text
        
        if not pymupdf_text.strip():
            return pdfminer_text
        
        merged = pdfminer_text
        pdfminer_lines = set(line.strip() for line in pdfminer_text.split('\n') if line.strip())
        pymupdf_lines = [line.strip() for line in pymupdf_text.split('\n') if line.strip()]
        
        unique_from_pymupdf = [line for line in pymupdf_lines if line not in pdfminer_lines]
        
        if unique_from_pymupdf:
            merged += "\n[Additional content from PDF structure analysis]\n"
            merged += "\n".join(unique_from_pymupdf)
        
        logger.info(f"Merged PDF extractions: pdfminer primary + {len(unique_from_pymupdf)} unique lines")
        return merged
    
    # ==================== DOCX Complete Extraction ====================
    
    def _extract_from_docx_complete(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text, hyperlinks, and images from DOCX"""
        extraction_details = {}
        
        # Extract images from DOCX with Base64 encoding
        logger.info("Extracting images from DOCX...")
        images = self._extract_images_from_docx(file_path)
        extraction_details["images"] = images
        logger.info(f"✅ Found {len(images)} image(s)")
        
        # Extract hyperlinks from DOCX
        logger.info("Extracting hyperlinks from DOCX...")
        hyperlinks = self._extract_hyperlinks_from_docx(file_path)
        extraction_details["hyperlinks"] = hyperlinks
        logger.info(f"✅ Found {len(hyperlinks)} hyperlink(s)")
        
        # Extract text from DOCX
        try:
            logger.info("Extracting text from DOCX...")
            text = docx2txt.process(file_path)
            extraction_details["docx"] = text
            logger.info("✅ DOCX extraction successful")
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            extraction_details["docx"] = ""
            text = ""
        
        if not text.strip():
            raise ResumeParsingError("DOCX extracted no text content")
        
        logger.info("DOCX extraction (complete) successful")
        return text, extraction_details
    
    def _extract_images_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract images from DOCX with Base64 encoding
        
        Returns list of dicts with image data
        """
        images = []
        
        try:
            doc = DocxDocument(file_path)
            
            image_index = 0
            
            # Extract from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        
                        # Get image format
                        content_type = image_part.content_type
                        image_format = content_type.split('/')[-1]  # e.g., 'jpeg' from 'image/jpeg'
                        
                        # Convert to Base64
                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                        
                        image_index += 1
                        images.append({
                            "image_name": f"docx_image_{image_index}.{image_format}",
                            "image_base64": image_base64,
                            "image_format": image_format,
                            "image_size": len(image_bytes),
                            "extraction_method": "python-docx"
                        })
                        
                        logger.debug(f"Extracted image from DOCX: {len(image_bytes)} bytes")
                    
                    except Exception as e:
                        logger.warning(f"Error extracting image from DOCX: {e}")
                        continue
            
            return images
        
        except Exception as e:
            logger.error(f"DOCX image extraction failed: {e}")
            return []
    
    def _extract_hyperlinks_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract hyperlinks from DOCX"""
        hyperlinks = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract hyperlinks from document relationships
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    try:
                        link_url = rel.target_ref
                        
                        # Find the link text by searching through paragraphs and runs
                        link_text = self._find_hyperlink_text(doc, rel.rId)
                        
                        hyperlinks.append({
                            "link_text": link_text,
                            "link_url": link_url,
                            "extraction_method": "python-docx"
                        })
                        
                        logger.debug(f"Found hyperlink in DOCX: {link_text} -> {link_url}")
                    
                    except Exception as e:
                        logger.warning(f"Error processing hyperlink: {e}")
                        continue
            
            return hyperlinks
        
        except Exception as e:
            logger.error(f"DOCX hyperlink extraction failed: {e}")
            return []
    
    def _find_hyperlink_text(self, doc, rId: str) -> str:
        """Find text associated with a hyperlink"""
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Check if this run contains the hyperlink
                    element = run._element
                    hyperlink_element = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink', 
                                                    namespaces=None)
                    if hyperlink_element is not None:
                        # Check if this is the hyperlink we're looking for
                        if hyperlink_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id') == rId:
                            return run.text
        except Exception as e:
            logger.debug(f"Error finding hyperlink text: {e}")
        
        return ""
    
    # ==================== Entity Extraction (spaCy) ====================
    
    def extract_entities_spacy(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities using spaCy"""
        try:
            doc = self.nlp(text[:1000000])
            
            entities_dict = {
                "PERSON": [],
                "ORG": [],
                "GPE": [],
                "DATE": [],
                "EVENT": [],
                "PRODUCT": [],
                "OTHER": []
            }
            
            for ent in doc.ents:
                label = ent.label_
                if label in entities_dict:
                    if ent.text not in entities_dict[label]:
                        entities_dict[label].append(ent.text)
                else:
                    if ent.text not in entities_dict["OTHER"]:
                        entities_dict["OTHER"].append(ent.text)
            
            logger.info("spaCy entity extraction completed")
            return entities_dict
        
        except Exception as e:
            logger.error(f"spaCy extraction failed: {e}")
            return {}
    
    # ==================== Regex Pattern Matching ====================
    
    def extract_patterns(self, text: str) -> Dict[str, Any]:
        """Extract information using regex patterns"""
        patterns = {
            "emails": self._extract_emails(text),
            "phone_numbers": self._extract_phone_numbers(text),
            "skills": self._extract_skills(text),
            "education": self._extract_education(text),
            "experience": self._extract_experience(text),
            "certifications": self._extract_certifications(text),
            "links": self._extract_links(text),
        }
        
        logger.info("Regex pattern extraction completed")
        return patterns
    
    def _extract_emails(self, text: str) -> List[str]:
        """Extract email addresses"""
        try:
            pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
            emails = list(set(re.findall(pattern, text)))
            return emails
        except Exception as e:
            logger.warning(f"Email extraction error: {e}")
            return []
    
    def _extract_phone_numbers(self, text: str) -> List[str]:
        """Extract phone numbers"""
        try:
            patterns = [
                r'\+?1?\d{9,15}',
                r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            ]
            phones = []
            for pattern in patterns:
                try:
                    phones.extend(re.findall(pattern, text))
                except Exception as e:
                    logger.warning(f"Phone pattern error: {e}")
            return list(set(phones))
        except Exception as e:
            logger.warning(f"Phone extraction error: {e}")
            return []
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills"""
        try:
            skills_keywords = [
                'Python', 'JavaScript', 'Java', 'C#', 'CSharp', 'Go', 'Rust', 'TypeScript', 'Kotlin',
                'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'FastAPI',
                'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git', 'GitLab', 'GitHub',
                'SQL', 'MongoDB', 'PostgreSQL', 'MySQL', 'Redis', 'Firebase',
                'Machine Learning', 'Deep Learning', 'NLP', 'Computer Vision', 'TensorFlow', 'PyTorch',
                'Data Science', 'Analytics', 'Tableau', 'Power BI', 'Pandas', 'NumPy',
                'REST API', 'GraphQL', 'Microservices', 'CI/CD', 'Agile', 'Scrum',
                'HVAC', 'Electrical', 'Plumbing', 'Mechanical', 'Construction'
            ]
            
            found_skills = []
            for skill in skills_keywords:
                try:
                    if re.search(re.escape(skill), text, re.IGNORECASE):
                        found_skills.append(skill)
                except Exception as e:
                    logger.warning(f"Skill pattern error for '{skill}': {e}")
            
            return list(set(found_skills))
        except Exception as e:
            logger.warning(f"Skills extraction error: {e}")
            return []
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education details"""
        try:
            education = []
            degrees = ['B.A.', 'B.S.', 'B.Tech', 'M.A.', 'M.S.', 'M.Tech', 'MBA', 'PhD', 'Ph.D', 'Bachelor', 'Master']
            
            for degree in degrees:
                try:
                    escaped_degree = re.escape(degree)
                    pattern = escaped_degree + r'.*?(?:in|from)?\s+([A-Za-z\s]+?)(?:from|at)?\s+([A-Za-z\s]+?)(?:\(|,|\.|\n|$)'
                    
                    matches = re.finditer(pattern, text, re.IGNORECASE)
                    for match in matches:
                        try:
                            education.append({
                                "degree": degree,
                                "field": match.group(1).strip() if match.group(1) else "",
                                "institution": match.group(2).strip() if match.group(2) else ""
                            })
                        except Exception as e:
                            logger.warning(f"Education match error: {e}")
                except Exception as e:
                    logger.warning(f"Education pattern error for '{degree}': {e}")
            
            return education
        except Exception as e:
            logger.warning(f"Education extraction error: {e}")
            return []
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience"""
        try:
            experience = []
            
            pattern = r'(\w+\s+(?:Engineer|Developer|Manager|Analyst|Designer|Architect|Lead|Senior|Junior|Specialist|Coordinator|Associate|Director|VP|Officer)).*?(?:at|@|-)?\s*([A-Z][A-Za-z\s&]+?)(?:,|\.|\(|\n|$)'
            
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    experience.append({
                        "job_title": match.group(1).strip(),
                        "company": match.group(2).strip() if match.group(2) else "",
                    })
                except Exception as e:
                    logger.warning(f"Experience match error: {e}")
            
            return experience
        except Exception as e:
            logger.warning(f"Experience extraction error: {e}")
            return []
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications and credentials"""
        try:
            certifications = []
            cert_keywords = [
                'AWS', 'GCP', 'Azure', 'Kubernetes', 'Docker', 'Certified',
                'PMP', 'CISISCO', 'SCRUM', 'Agile', 'GCP Certified', 'AWS Certified',
                'Azure Certified', 'CompTIA', 'Security+', 'CCNA', 'LPIC'
            ]
            
            for cert in cert_keywords:
                try:
                    if re.search(re.escape(cert), text, re.IGNORECASE):
                        certifications.append(cert)
                except Exception as e:
                    logger.warning(f"Certification pattern error for '{cert}': {e}")
            
            return list(set(certifications))
        except Exception as e:
            logger.warning(f"Certifications extraction error: {e}")
            return []
    
    def _extract_links(self, text: str) -> List[str]:
        """Extract URLs from text"""
        try:
            pattern = r'https?://[^\s]+|www\.[^\s]+'
            links = list(set(re.findall(pattern, text)))
            return links
        except Exception as e:
            logger.warning(f"Links extraction error: {e}")
            return []
    
    # ==================== Merge & Deduplicate ====================
    
    def merge_extraction_results(
        self,
        spacy_entities: Dict[str, List[str]],
        regex_patterns: Dict[str, Any],
        hyperlinks: List[Dict[str, Any]],
        images: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Merge all extraction results"""
        merged = {
            "personal_info": {
                "names": list(set(spacy_entities.get("PERSON", []))),
                "emails": regex_patterns.get("emails", []),
                "phone_numbers": regex_patterns.get("phone_numbers", []),
                "links": regex_patterns.get("links", []),
            },
            "professional_info": {
                "organizations": list(set(spacy_entities.get("ORG", []))),
                "locations": list(set(spacy_entities.get("GPE", []))),
                "job_titles": [
                    exp["job_title"] for exp in regex_patterns.get("experience", [])
                ],
                "companies": [
                    exp["company"] for exp in regex_patterns.get("experience", [])
                    if exp["company"]
                ],
            },
            "skills_and_expertise": {
                "technical_skills": regex_patterns.get("skills", []),
                "certifications": regex_patterns.get("certifications", []),
            },
            "education": regex_patterns.get("education", []),
            "experience": regex_patterns.get("experience", []),
            "hyperlinks": hyperlinks,
            "images": images,  # NEW: Separate images key
            "metadata": {
                "extraction_timestamp": datetime.now().isoformat(),
                "extraction_methods": [
                    "Dual PDF (PyMuPDF + pdfminer.six)",
                    "spaCy NER",
                    "Regex Patterns",
                    "Hyperlink Extraction",
                    "Image Extraction (Base64)"
                ],
            }
        }
        
        return merged
    
    # ==================== Main Parse Method ====================
    
    def parse(self, file_path: str, include_raw_text: bool = False) -> Dict[str, Any]:
        """Complete resume parsing with all features
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            include_raw_text: Whether to include the raw text in the response (can be large)
            
        Returns:
            Dict containing parsed resume data, hyperlinks, and images
        """
        try:
            logger.info(f"Starting resume parsing for: {file_path}")
            
            # Extract all content
            raw_text, extraction_details = self.extract_resume_text(file_path)
            hyperlinks = extraction_details.get("hyperlinks", [])
            images = extraction_details.get("images", [])
            
            # Extract entities
            spacy_entities = self.extract_entities_spacy(raw_text)
            
            # Extract patterns
            regex_patterns = self.extract_patterns(raw_text)
            
            # Merge all results
            merged_result = self.merge_extraction_results(
                spacy_entities,
                regex_patterns,
                hyperlinks,
                images
            )
            
            # Add metadata
            if include_raw_text:
                merged_result["raw_text"] = raw_text
            merged_result["file_path"] = str(file_path)
            merged_result["status"] = "success"
            
            # Add extraction details
            merged_result["metadata"] = {
                "extraction_date": datetime.now().isoformat(),
                "hyperlinks_count": len(hyperlinks),
                "images_count": len(images),
                "file_size": Path(file_path).stat().st_size if Path(file_path).exists() else 0
            }
            
            logger.info(
                f"Resume parsing completed: {len(hyperlinks)} hyperlinks, "
                f"{len(images)} images found"
            )
            return merged_result
        
        except ResumeParsingError as e:
            logger.error(f"Resume parsing failed: {e}")
            return {
                "status": "error",
                "error_message": str(e),
                "file_path": str(file_path),
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Unexpected error during parsing: {e}")
            return {
                "status": "error",
                "error_message": f"Unexpected error: {str(e)}",
                "file_path": str(file_path),
                "timestamp": datetime.now().isoformat()
            }

# ==================== Usage Example ====================

if __name__ == "__main__":
    import sys
    import json
    
    if len(sys.argv) < 2:
        print("Usage: python pdf-with-image-parse.py <path_to_resume> [--include-raw-text]")
        sys.exit(1)
        
    file_path = sys.argv[1]
    
    try:
        parser = DualPDFDocxParserWithLinksAndImages()
        result = parser.parse(file_path, include_raw_text=True)
        
        # Print the JSON result to stdout
        print(json.dumps(result, indent=2, ensure_ascii=False))
        
    except Exception as e:
        error_result = {
            "status": "error",
            "error_message": f"Fatal error: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }
        print(json.dumps(error_result, indent=2, ensure_ascii=False))
        sys.exit(1)